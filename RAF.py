#!/usr/bin/env python
# coding: utf-8

# ## Download all the necessary libraries

# In[ ]:


import pandas as pd
import numpy as np
import easygui
import sys
import logging
from docx import Document
from docxtpl import DocxTemplate
from docx.shared import Cm, Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import lxml.etree
import os
import re
from pdf2image import convert_from_path
import decimal
import datetime
import time
from time import strptime
from pandas import Timestamp
from meteostat import Point, Daily
import shutil
import beepy as beep
import ctypes
import win32com.client
os.environ['PYGAME_HIDE_SUPPORT_PROMPT'] = "hide"
from pygame import mixer


# In[ ]:


# from IPython.core.display import display, HTML
# display(HTML("<style>.container { width:100% !important; }</style>"))
import warnings
warnings.simplefilter(action='ignore', category=FutureWarning)
pd.options.mode.chained_assignment = None  # default='warn'
pd.options.display.float_format = "{:,.2f}".format
logger = logging.getLogger('ftpuploader')


# In[ ]:


# Create function for sounds making
def beep(sound):
    mixer.init() 
    sound=mixer.Sound(os.getcwd() + '/SYSTEM/Sounds/{}.wav'.format(sound))
    sound.play()


# In[ ]:





# ## Download and preprocess source files

# In[ ]:


print('Uploading file "Estimate"...')


# In[ ]:


# Define the folder to work with
 # Add sound
beep(2)
fn = 'Work ' + easygui.buttonbox('Choose the Work folder', 'Attention!!!', ('1', '2', '3'))


# In[ ]:


try:
    # Initially define Smeta file name and script status
    smeta = 'Estimate'
    scrp = 'creation'
    # Search for Smeta files
    for smeta_file in os.listdir(os.getcwd() + '/{}'.format(fn)):
        if smeta_file.startswith('Estimate') and os.path.splitext(smeta_file)[0][-1].isdigit()==True:
            smeta = os.path.splitext(smeta_file)[0]
            scrp = 'correction'  
    
    # Check if the file is closed
    agfile = 'open'
    while agfile == 'open':
        try: 
            os.rename(os.getcwd() + '/{}/{}.xlsx'.format(fn, smeta), os.getcwd() + '/{}/Estimate OCcheck.xlsx'.format(fn))
            os.rename(os.getcwd() + '/{}/Estimate OCcheck.xlsx'.format(fn), os.getcwd() + '/{}/{}.xlsx'.format(fn, smeta))
            agfile = 'closed'
        except BaseException  as e:
            beep(4)
            easygui.msgbox('Close file "Estimate": \n' + str(e), title='Error!')
    
    # Read Agreement table
    ag = pd.read_excel(os.getcwd() + '/{}/{}.xlsx'.format(fn, smeta), sheet_name='Договор', usecols=[0, 1, 4], nrows=18)
    
    # Change research end's date format from Timestamp to Str
    if type(ag['Значение'].loc[3]) is datetime.datetime:
       ag['Значение'].loc[3] = datetime.datetime.strftime(ag['Значение'].loc[3], '%d.%m.%Y')
    
    # What if it is standard form
    if ag['Unnamed: 4'][3] == 'Стандартная':
        # Read Agreement Data table if it is exist
        for file in os.listdir(os.getcwd() + '/{}'.format(fn)):
            if file.startswith('Contract_data'):
                ad = pd.read_excel(os.getcwd() + '/{}/{}'.format(fn, file))
                # Replace ag parameters with ad
                ag['Значение'].loc[0] = ad['Значение'].loc[1]
                ag['Значение'].loc[1] = ad['Значение'].loc[2]
                ag['Значение'].loc[2] = ad['Значение'].loc[3]
                ag['Значение'].loc[3] = ad['Значение'].loc[4]
                ag['Значение'].loc[4] = ad['Значение'].loc[7]
                ag['Значение'].loc[7] = ad['Значение'].loc[9]
        # Define interesant
        interesant = ''
    # What if it is not standard form
    else:
        ag['Значение'].loc[4] = ag['Значение'].loc[5]
        ag['Значение'].loc[7] = 'ИП Трубников Алексей Владимирович'
        # Define interesant
        interesant = '(в интересах ' + ag['Значение'].loc[8] +')'
        
    # Create outside temperature variable
    ostemp = ag['Значение'].loc[15]
        
except BaseException  as e:
    beep(4)
    easygui.msgbox('Failed to upload tab "Договор": \n' + str(e), title='Error!')
    logger.error(str(e))
    sys.exit()


# In[ ]:





# In[ ]:


try:
    # Read Teplovisor table
    tv = pd.read_excel(os.getcwd() + '/{}/{}.xlsx'.format(fn, smeta), sheet_name='Тепловизор', usecols=[0, 1, 2, 3, 4])
    # Round the last column
    tv['Отклонение температуры, °С'] = tv['Отклонение температуры, °С'].round(1)
    # Delete NaNs
    tv = tv.dropna()
    # Round to 0 decimals
    tv[tv.columns[1]] = tv[tv.columns[1]].values.astype(np.int64)
    # Convert data to str format
    tv = tv.astype(str)
    
except BaseException  as e:
    beep(4)
    easygui.msgbox('Failed to upload tab "Тепловизор": \n' + str(e), title='Error!')
    logger.error(str(e))
    sys.exit()


# In[ ]:





# In[ ]:


try:
    # Read Defects table
    sd = pd.read_excel(os.getcwd() + '/{}/{}.xlsx'.format(fn, smeta), sheet_name='Перечень дефектов', usecols=[0, 1, 2, 3])
    # Delete unnecessary NaNs
    # Delete rows wo foto
    for i in range(len(sd)):
        if pd.isna(sd['Фото'].loc[i]) == True and pd.isna(sd['№ п/п'].loc[i]) == False:
            sd.drop(index = i, inplace = True)
    sd = sd.reset_index(drop=True)
    # Delete subtitle rows wo content
    for i in range(1, len(sd)):
        if pd.isna(sd['№ п/п'].loc[i-1]) == True and pd.isna(sd['№ п/п'].loc[i]) == True:
            sd.drop(index = i-1, inplace = True)
    sd = sd.reset_index(drop=True)        
    # Delete the last row wo content
    if pd.isna(sd['№ п/п'].loc[len(sd)-1]) == True:
        sd.drop(index = len(sd)-1, inplace = True)
    # Replace ''№ п/п' with string simbol to delete decimals
    for i in range(len(sd)):
        if pd.isna(sd['№ п/п'].loc[i]) == False:
            sd['№ п/п'].loc[i] = 'a'
    # Renumerate rows within the table
    for i in range(1, len(sd)):
        if pd.isna(sd['№ п/п'].loc[i-1]) == True and pd.isna(sd['№ п/п'].loc[i]) == False:
            sd['№ п/п'].loc[i] = 1
        if pd.isna(sd['№ п/п'].loc[i-1]) == False and pd.isna(sd['№ п/п'].loc[i]) == False:
            sd['№ п/п'].loc[i] = int(sd['№ п/п'].loc[i-1]+1)
    # Replace NaNs with blanks
    sd.fillna('', inplace=True)
    # Make photos numbers readable 
    for i in range(len(sd)):
        sd['Фото'][i] = re.sub('/$', 'мм)', str(sd['Фото'][i]))
        sd['Фото'][i] = str(sd['Фото'][i]).replace('/+', 'мм)\n').replace('/', '(').replace('+', '\n')

except BaseException  as e:
    beep(4)
    easygui.msgbox('Failed to upload tab "Дефекты": \n' + str(e), title='Error!')
    logger.error(str(e))
    sys.exit()


# In[ ]:





# In[ ]:


try:
    # Read Vedomost table
    dv = pd.read_excel(os.getcwd() + '/{}/{}.xlsx'.format(fn, smeta), sheet_name='Дефектная ведомость')
    # Delete last 7 columns
    dv = dv.iloc[:, :-7]
    # Replace NaNs with blanks
    dv.fillna('', inplace=True)

    # Delete unnecessary NaNs and 0's
    # Delete 0's
    for i in range(len(dv)):
        if dv['Всего'].loc[i] == 0:
            dv.drop(index = i, inplace = True)
    dv = dv.reset_index(drop=True)
    # Delete last 2 rows with NaNs
    dv.drop(dv.tail(2).index,  inplace = True)
    # Renumerate rows within the table
    dv['№ п/п'] = dv.index

    # Round it
    for column in dv.columns[3:]:
        for i in range(1, len(dv)):
            # Replace 0 with ''
            if  dv[column].loc[i] == 0:
                dv[column].loc[i] = ''
            if  dv[column].loc[i] != '' and dv['Ед. изм.'].loc[i] in ['м', 'м2']:
                # Round to 2 decimals
                dv[column].loc[i] = round(dv[column].loc[i], 2)
            elif dv[column].loc[i] != '':
                # Round to 0 decimals
                dv[column].loc[i] = int(dv[column].loc[i])

except BaseException  as e:
    beep(4)
    easygui.msgbox('Failed to upload tab "Ведомость": \n' + str(e), title='Error!')
    logger.error(str(e))
    sys.exit()


# In[ ]:





# In[ ]:


try:
    # Read Prices and Contractors tables
    pr = pd.read_excel(os.getcwd() + '/{}/{}.xlsx'.format(fn, smeta), usecols = [0,1,2,3,4,5,6,7,8,9], sheet_name='Расчеты')
    # Find where to end
    end = pr[pr['Наименование работ и материалов'] == 'Итого работы и материалы'].index[0] + 1
    # Read Prices table
    pr = pd.read_excel(os.getcwd() + '/{}/{}.xlsx'.format(fn, smeta), nrows = end, usecols = [0,1,2,3,4,5,6,7,8,9], sheet_name='Расчеты')
    # Delete 0's
    for i in range(len(pr)):
        if pr['Общая стоимость, руб.'].loc[i] == 0:
            pr.drop(index = i, inplace = True)
    pr = pr.reset_index(drop=True)
    # Renumerate rows within the table
    pr['№ п/п'].loc[1] = 1
    for i in range(2, len(pr)):
        if pd.isna(pr['№ п/п'].loc[i-1]) == True and pd.isna(pr['№ п/п'].loc[i]) == False:
            pr['№ п/п'].loc[i] = 1
        if pd.isna(pr['№ п/п'].loc[i-1]) == False and pd.isna(pr['№ п/п'].loc[i]) == False:
            pr['№ п/п'].loc[i] = int(pr['№ п/п'].loc[i-1]+1)
    # Replace the rest NaNs with blanks
    pr.fillna('', inplace=True)
    # Round to 2 decimals
    for column in pr.columns[3:10]:
        for i in range(1, len(pr)):
            if pr[column].loc[i] != '':
                pr[column].loc[i] = round(pr[column].loc[i], 2)
    # Round to 0 decimals
     # Prices
    for column in pr.columns[4:8]:
        for i in range(0, len(pr)):
            if pr[column].loc[i] != '':
                pr[column].loc[i] = int(pr[column].loc[i])
     # Quantities and Numeration
    for i in range(1, len(pr)):
        if  pr['Ед. изм.'].loc[i] != '' and pr['Ед. изм.'].loc[i] not in ['м', 'м2']:
            pr['Кол-во'].loc[i] = int(pr['Кол-во'].loc[i])
        if  pr['№ п/п'].loc[i] != '':
            pr['№ п/п'].loc[i] = int(pr['№ п/п'].loc[i])

except BaseException  as e:
    beep(4)
    easygui.msgbox('Failed to upload tab "Стоимость": \n' + str(e), title='Error!')
    logger.error(str(e))
    sys.exit()


# In[ ]:





# In[ ]:


try:
    # Read Windows table
    wn = pd.read_excel(os.getcwd() + '/{}/{}.xlsx'.format(fn, smeta), usecols = [0,1,2,3,4,5,6,7], sheet_name='Окна')
    # Delete last 2 rows with NaNs
    wn.drop(wn.tail(2).index,  inplace = True)
    # Check if there are standard windows
    standwind = wn['Кол-во'].loc[1]
    # Delete 0's
    for i in range(len(wn)):
        if wn['Общая стоимость, руб.'].loc[i] == 0:
            wn.drop(index = i, inplace = True)
    wn = wn.reset_index(drop=True)
    # Renumerate rows within the tablee
    wn['№ п/п'].loc[1] = 1
    for i in range(2, len(wn)-1):
        wn['№ п/п'].loc[i] = int(wn['№ п/п'].loc[i-1]+1)
    # Replace the rest NaNs with blanks
    wn.fillna('', inplace=True)    
    # Round to 2 decimals
    for column in wn.columns[[3, 6, 7]]:
        for i in range(1, len(wn)):
            if wn[column].loc[i] != '':
                wn[column].loc[i] = round(wn[column].loc[i], 2)
    # Round to 0 decimals
     # Prices
    for column in wn.columns[4:6]:
        for i in range(0, len(wn)-1):
            wn[column].loc[i] = int(wn[column].loc[i])
     # Quantities and Numeration 
    for i in range(1, len(wn)-1):
        if  wn['Ед. изм.'].loc[i] == 'шт.':
            wn['Кол-во'].loc[i] = int(wn['Кол-во'].loc[i])
        wn['№ п/п'].loc[i] = int(wn['№ п/п'].loc[i])

except BaseException  as e:
    beep(4)
    easygui.msgbox('Failed to upload tab "Окна": \n' + str(e), title='Error!')
    logger.error(str(e))
    sys.exit()


# In[ ]:





# In[ ]:


try:
    # Read Ceiling table
    cl = pd.read_excel(os.getcwd() + '/{}/{}.xlsx'.format(fn, smeta), usecols = [0,1,2,3,4,5,6,7,8], sheet_name='Потолок')
#     # Drop the last two rows
#     cl.drop(cl.tail(2).index, inplace = True)
    # Delete 0's
    for i in range(len(cl)):
        if cl['Общая стоимость, руб.'].loc[i] == 0:
            cl.drop(index = i, inplace = True)
    cl = cl.reset_index(drop=True)
    # Renumerate rows within the tablee
    cl['№ п/п'].loc[1] = 1
    for i in range(2, len(cl)-1):
        cl['№ п/п'].loc[i] = int(cl['№ п/п'].loc[i-1]+1)
    # Replace the rest NaNs with blanks
    cl.fillna('', inplace=True) 
    # Round to 2 decimals
    for column in cl.columns[[3, 7, 8]]:
        for i in range(1, len(cl)):
            if cl[column].loc[i] != '':
                cl[column].loc[i] = round(cl[column].loc[i], 2)
    # Round to 0 decimals
     # Prices
    for column in cl.columns[4:7]:
        for i in range(0, len(cl)-1):
            cl[column].loc[i] = int(cl[column].loc[i])
     # Quantities and Numeration
    for i in range(1, len(cl)-1):
        if  cl['Ед. изм.'].loc[i] == 'шт.':
            cl['Кол-во'].loc[i] = int(cl['Кол-во'].loc[i])
        cl['№ п/п'].loc[i] = int(cl['№ п/п'].loc[i])

except BaseException  as e:
    beep(4)
    easygui.msgbox('Failed to upload tab "Потолок": \n' + str(e), title='Error!')
    logger.error(str(e))
    sys.exit()


# In[ ]:





# In[ ]:


try:
    # Read Contractors table
    cn = pd.read_excel(os.getcwd() + '/{}/{}.xlsx'.format(fn, smeta), sheet_name='Ресурсы')

    # What if there are no windows job
    if wn.shape[0]==1:
        # Delete windows contractors
        cn.drop(index = [4, 5], inplace = True)

    # What if there are no ceiling job
    if cl.shape[0]==1:
        # Delete ceiling contractors
        cn.drop(index = [6, 7, 8], inplace = True)

    # Reset index
    cn = cn.reset_index(drop=True)
    # Renumerate rows within the table
    cn['№'] = cn.index + 1   

    # Create dataframe to make contracrtors Table from it
    cnt = cn.drop('Наименование ресурса', 1)

    # Create dataframe to make contracrtors Listliketable from it
    cnl = pd.DataFrame(columns=['№', 'Name'])
    cnl['№'] = cn['№']
    for i in range(len(cnl)):
        cnl['Name'].loc[i] = cn['Наименование ресурса'].astype(str).loc[i]+' ('+cn['Адрес ресурса'].astype(str).loc[i]+')'

except BaseException  as e:
    beep(4)
    easygui.msgbox('Failed to upload tab "Ресурсы": \n' + str(e), title='Error!')
    logger.error(str(e))
    sys.exit()


# In[ ]:





# In[ ]:


try:
    # Read Certificates table
    ct = pd.read_excel(os.getcwd() + '/Source/Verifications/Equipment.xlsx')
    # Get the revision date and convert it to timestamp
    rd = Timestamp(np.datetime64(datetime.datetime.strptime(re.sub(r'\s\w\.+$', '', ag['Значение'][2]), "%d.%m.%Y")))
    # Filter only actual certificates
     # Filter Teplovisor certificates
    if  tv.shape[0] == 0:
        for i in range(len(ct)):
            if ct['Оборудование'].loc[i].startswith('Тепловизор') | ct['Оборудование'].loc[i].startswith('Гигрометр'):
                ct.drop(index = i, inplace = True)         
     # Filter actual dates
    ct = ct[((ct['Срок действия поверки'] >= rd) & (ct['Дата поверки'] <= rd)) | (ct['Документ'].isna())]
    # Delete duplicates
    ct = ct[ct['Оборудование'].duplicated() == False]
    ct = ct.reset_index(drop=True)
    # Save documents as anothe table
    ctd = ct['№ документа']
    # Renumerate rows within the table
    ct['№ п/п'] = ct.index + 1
    # Change date format
    ct['Дата поверки'] = pd.to_datetime(ct['Дата поверки']).dt.strftime('%d.%m.%Y')
    # Replace the rest NaNs with blanks
    ct.fillna('', inplace=True)
    # Combine Columns to one row
    ct['Description'] = ''
    for i in range(len(ct)):
        if ct['Документ'].loc[i] == '':
            ct['Description'].loc[i] = ct['№ п/п'].astype(str).loc[i]+'. '+ct['Оборудование'].loc[i]
        else:
            ct['Description'].loc[i] = ct['№ п/п'].astype(str).loc[i]+'. '+ct['Оборудование'].astype(str).loc[i]+' '+ct['Документ'].astype(str).loc[i]+' №' +ct['№ документа'].astype(str).loc[i]+' от '+ct['Дата поверки'].astype(str).loc[i]+'г.'
    ct = ct['Description']

except BaseException  as e:
    beep(4)
    easygui.msgbox('Failed to upload tab "Equipment": \n' + str(e), title='Error!')
    logger.error(str(e))
    sys.exit()


# In[ ]:


print('  Done')


# In[ ]:





# ## Create all the necessary functions

# In[ ]:


# Create function for deleting paragraphs
def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


# In[ ]:


# Create function to move tables
def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


# In[ ]:


# Create function to repeat tables headers
def set_repeat_table_header(row):
    """ set repeat table row on every new page
    """
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)
    return row


# In[ ]:


# Create function для создания числа прописью
units = (
    u'ноль',
    (u'один', u'одна'),
    (u'два', u'две'),
    u'три', u'четыре', u'пять',
    u'шесть', u'семь', u'восемь', u'девять'
)
teens = (
    u'десять', u'одиннадцать',
    u'двенадцать', u'тринадцать',
    u'четырнадцать', u'пятнадцать',
    u'шестнадцать', u'семнадцать',
    u'восемнадцать', u'девятнадцать'
)
tens = (
    teens,
    u'двадцать', u'тридцать',
    u'сорок', u'пятьдесят',
    u'шестьдесят', u'семьдесят',
    u'восемьдесят', u'девяносто'
)
hundreds = (
    u'сто', u'двести',
    u'триста', u'четыреста',
    u'пятьсот', u'шестьсот',
    u'семьсот', u'восемьсот',
    u'девятьсот'
)
orders = (# plural forms and gender
    #((u'', u'', u''), 'm'), # ((u'рубль', u'рубля', u'рублей'), 'm'), # ((u'копейка', u'копейки', u'копеек'), 'f')
    ((u'тысяча', u'тысячи', u'тысяч'), 'f'),
    ((u'миллион', u'миллиона', u'миллионов'), 'm'),
    ((u'миллиард', u'миллиарда', u'миллиардов'), 'm'),
)
minus = u'минус'

def thousand(rest, sex):
    """Converts numbers from 19 to 999"""
    prev = 0
    plural = 2
    name = []
    use_teens = rest % 100 >= 10 and rest % 100 <= 19
    if not use_teens:
        data = ((units, 10), (tens, 100), (hundreds, 1000))
    else:
        data = ((teens, 10), (hundreds, 1000))
    for names, x in data:
        cur = int(((rest - prev) % x) * 10 / x)
        prev = rest % x
        if x == 10 and use_teens:
            plural = 2
            name.append(teens[cur])
        elif cur == 0:
            continue
        elif x == 10:
            name_ = names[cur]
            if isinstance(name_, tuple):
                name_ = name_[0 if sex == 'm' else 1]
            name.append(name_)
            if cur >= 2 and cur <= 4:
                plural = 1
            elif cur == 1:
                plural = 0
            else:
                plural = 2
        else:
            name.append(names[cur-1])
    return plural, name

def num2text(num, main_units=((u'', u'', u''), 'm')):
    """
    http://ru.wikipedia.org/wiki/Gettext#.D0.9C.D0.BD.D0.BE.D0.B6.D0.B5.D1.81.\
    D1.82.D0.B2.D0.B5.D0.BD.D0.BD.D1.8B.D0.B5_.D1.87.D0.B8.D1.81.D0.BB.D0.B0_2
    """
    _orders = (main_units,) + orders
    if num == 0:
        return ' '.join((units[0], _orders[0][0][2])).strip() # ноль
    rest = abs(num)
    ord = 0
    name = []
    while rest > 0:
        plural, nme = thousand(rest % 1000, _orders[ord][1])
        if nme or ord == 0:
            name.append(_orders[ord][0][plural])
        name += nme
        rest = int(rest / 1000)
        ord += 1
    if num < 0:
        name.append(minus)
    name.reverse()
    return ' '.join(name).strip()


# In[ ]:





# ## Agreement

# In[ ]:


print('Uploading template "Сonclusions"...')


# In[ ]:


# Get the name of the folder and conclusion
 # Get apartment number
apnr = re.findall(r'\d+', ag['Значение'][4])[-1]
 # Get contract number
agnr = ag['Значение'][0][:3]
if ag['Unnamed: 4'][0] == 'Другая':
    name = ag['Unnamed: 4'][1] + '_' + apnr
else:
    name = ag['Unnamed: 4'][0] + '_' + apnr


# In[ ]:


# Count research start time
# Replace - with : and convert to datetime
ag['Значение'][10] = ag['Значение'][10].replace('-', ':')
# Get Hours and minutes
(h, m) = ag['Значение'][10].split(':')
# Find how much time past research start
delta = 5 + int(m) % 5
# Count research start time
resstart = datetime.timedelta(hours=int(h), minutes=int(m)) - datetime.timedelta(minutes=delta)
# Convert it to datetime
resstart = pd.to_datetime(str(resstart)).strftime('%H:%M')


# In[ ]:


# Count research end time
# Replace - with : and convert to datetime
ag['Значение'][11] = ag['Значение'][11].replace('-', ':')
# Get Hours and minutes
(h, m) = ag['Значение'][11].split(':')
# Find how much time before research end
delta = 10 - int(m) % 5
# Count research end time
resend = datetime.timedelta(hours=int(h), minutes=int(m)) + datetime.timedelta(minutes=delta)
# Convert it to datetime
resend = pd.to_datetime(str(resend)).strftime('%H:%M')


# In[ ]:


try:
    # Download Conclusion template
    c = DocxTemplate(os.getcwd() + '/Source/Templates/Conclusion.docx')

except BaseException  as e:
    beep(4)
    easygui.msgbox('Failed to upload template "Conclusion": \n' + str(e), title='Error!')
    logger.error(str(e))
    sys.exit()


# In[ ]:


# Delete unnecessary paragraph and add signature to the standard conclusion
if ag['Unnamed: 4'][3] == 'Стандартная':
    # Delete unnecessary paragraph
     # Find the paragraph
    indx=0
    for p in c.paragraphs:
        if p.text.startswith('По характеру отклонений дефектов, экспертизой было установлено'):
            trpindx = indx
        indx+=1
     # Delete it
    delete_paragraph(c.paragraphs[trpindx])
    
    # Add signature
     # Find where to add the signature
    indx=0
    for p in c.tables[0].cell(0,0).paragraphs:
        if p.text.startswith('Рабинович А.В'):
            sndpindx = indx
        indx+=1
     # Add it
    c.tables[0].cell(0,0).paragraphs[sndpindx+1].add_run().add_picture(os.getcwd() + '/Source/Templates/Signature.png', height=Cm(3))


# In[ ]:


print('  Done')


# In[ ]:





# ## Styles

# In[ ]:


## Create all the necessary styles


# In[ ]:


# Create new font style for tables content
cstyles = c.styles
tablestyle = cstyles.add_style('TableStyle', WD_STYLE_TYPE.PARAGRAPH)
font = tablestyle.font
font.name = 'Calibri'
font.size = Pt(10)


# In[ ]:


# Create new font style for tables subtitles
subtitlestyle = cstyles.add_style('SubtitleStyle', WD_STYLE_TYPE.PARAGRAPH)
font = subtitlestyle.font
font.name = 'Calibri'
font.bold= True
font.size = Pt(10)


# In[ ]:





# ## Teplovisor

# In[ ]:


print('Filling Thermal Imager readings...')


# In[ ]:


# What if there is no teplovisor research:
if  tv.shape[0] == 0:
    
    # Find Teplovisor GOST paragraphs
    indx=0
    for p in c.paragraphs[0:99]:
        if re.search(r'ГОСТ Р 54852-2011 «Здания и сооружения. Метод тепловизионного ', p.text) is not None:
            gostindx = indx
        indx+=1
    
    # Delete Teplovisor GOST paragraphs
    delete_paragraph(c.paragraphs[gostindx])
    
    # Find Teplovisor research start and end paragraphs
    indx=0
    for p in c.paragraphs:
        if p.text == 'Обследование тепловой защиты':
            startindx = indx
        elif p.text == 'Обследование помещения':
            endindx = indx
        indx+=1

    # Define Teplovisor paragraphs
    tp = c.paragraphs[startindx:endindx]
    
    # Delete Teplovisor paragraphs
    for i in tp:
        delete_paragraph(i)
        
    # Find Teplovisor norms table
    indx=0
    for t in c.tables:
        if t.cell(0,0).paragraphs[0].text.startswith('СП 50.13330.2012'):
            tindx = indx
        indx+=1
    
    # Delete Teplovisor norms table
    c.tables[1]._element.getparent().remove(c.tables[tindx]._element)
    
# What if there is teplovisor research:
else:
    
    # Get outside temperature on the research date
    if ostemp == '-':
        try:
            # Get research date in datetime format
            rddt = datetime.datetime.strptime(re.sub(r'\s\w\.+$', '', ag['Значение'][2]), "%d.%m.%Y")

            # Create Location Point for Moscow
            location = Point(55.7512, 37.6184)

            # Get daily data
            data = Daily(location, rddt, rddt)
            data = data.fetch()

            # Get only temperature
            ostemp = int(round(data['tavg'][0], 0))

        except BaseException  as e:
            beep(4)
            ostemp = '-'
            ostemp = easygui.enterbox('Failed to upload outdoor temperature, enter it manually:')
    
    # Convert DataFrame to Table and add it to the Conclusion
    # add a table to the end and create a reference variable
    t = c.add_table(tv.shape[0]+1, tv.shape[1])
    # add the header rows.
    for j in range(tv.shape[-1]):
        t.cell(0,j).text = tv.columns[j]
    # add the rest of the data frame
    for i in range(tv.shape[0]):
        for j in range(tv.shape[-1]):
            t.cell(i+1,j).text = str(tv.values[i,j])
    # Change table style
    t.style = 'Table Grid'
    # Align the table
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set table columns width
    widths = [Cm(1.0), Cm(3.0), Cm(6.0), Cm(3.0), Cm(3.0)]
    # Apply these widths to the table, center it's cells and change font to Pt10 (TableStyle)
    for row in t.rows:
        for idx, w in enumerate(widths):
            row.cells[idx].width = w
            row.cells[idx].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            row.cells[idx].paragraphs[0].style = c.styles['TableStyle']
            
    # Find the place where to add Teplovisor table
    indx=0
    for p in c.paragraphs:
        if p.text == 'Результаты обследования:':
            tvindx = indx
        indx+=1  
            
    # Move Teplovisor table to it's place
    move_table_after(t, c.paragraphs[tvindx])            


# In[ ]:


print('  Done')


# In[ ]:





# ## Defects

# In[ ]:


print('Filling Defect list...')


# In[ ]:


# Convert DataFrame to Table and add it to the Conclusion
# add a table to the end and create a reference variable
d = c.add_table(sd.shape[0]+1, sd.shape[1])
# add the header rows.
for j in range(sd.shape[-1]):
    d.cell(0,j).text = sd.columns[j]
# add the rest of the data frame
for i in range(sd.shape[0]):
    for j in range(sd.shape[-1]):
        d.cell(i+1,j).text = str(sd.values[i,j])
# Change table style
d.style = 'Table Grid'
# Align the table
d.alignment = WD_TABLE_ALIGNMENT.CENTER


# In[ ]:


# Avoid autofitting
d.autofit = False 
d.allow_autofit = False


# In[ ]:


# Set table columns width
widths = [Cm(0.9), Cm(4.2), Cm(3.3), Cm(10.4)]
# Apply these widths to the table, vertically center it's cells and change font to Pt10 (TableStyle)
for row in d.rows:
    for idx, w in enumerate(widths):
        row.cells[idx].width = w
#         row.cells[idx].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        row.cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row.cells[idx].paragraphs[0].style = c.styles['TableStyle']
# Center header and columns 0 and 2 horisontally
for idx, w in enumerate(widths):
    d.rows[0].cells[idx].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
for row in d.rows:
    for idx in [0, 2]:
        row.cells[idx].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER


# In[ ]:


# Merge subtitles cells 3 to 1 and delete 2 empty strings after merging
for row in d.rows:
    if row.cells[0].paragraphs[0].text == '':
        row.cells[1].merge(row.cells[2].merge(row.cells[3])).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        delete_paragraph(row.cells[1].paragraphs[-1])
        delete_paragraph(row.cells[1].paragraphs[-1])
        # Change merged cell stylle to bold
        row.cells[1].paragraphs[0].style = c.styles['SubtitleStyle']


# In[ ]:


# Repeat header rows
set_repeat_table_header(d.rows[0])


# In[ ]:


# Find the place where to add Defects table
indx=0
for p in c.paragraphs:
    if p.text == 'Обследование помещения':
        sdindx = indx
    indx+=1


# In[ ]:


# Move Defects table to it's place
move_table_after(d, c.paragraphs[sdindx])


# In[ ]:


print('  Done')


# In[ ]:





# ## The Walls

# In[ ]:


try:
    #Create the table where to put The Walls picture in
    twt = c.add_table(1, 1)
    # Add paragraph to the first cell and run it
    twp = twt.rows[0].cells[0].add_paragraph()
    twr = twp.add_run('На плане квартиры ниже красным цветом указаны стены, поверхность которых подлежит выравниванию.')
    # Add picture to this paragraph
    for filename in os.listdir(os.getcwd() + '/{}/'.format(fn)):
        if re.search(r'Walls', filename) is not None:
            twr.add_picture(os.getcwd() + '/{}/{}'.format(fn, filename), height=Cm(18))
    # Delete empty string before picture
    delete_paragraph(twt.rows[0].cells[0].paragraphs[-2])
        
except BaseException  as e:
    beep(4)
    easygui.msgbox('Failed to upload image "Walls": \n' + str(e), title='Error!')
    logger.error(str(e))
    sys.exit()        


# In[ ]:


# Move The Wall picture table to it's place
move_table_after(twt, c.paragraphs[sdindx+1])


# In[ ]:





# ## The Plan

# In[ ]:


try:
    # Create the table where to put The Plan picture in
    tpt = c.add_table(1, 1)
    # Add paragraph to the first cell and run it
    tpp = tpt.rows[0].cells[0].add_paragraph()
    tpr = tpp.add_run()
    # Add picture to this paragraph
    for filename in os.listdir(os.getcwd() + '/{}/'.format(fn)):
        if filename.endswith(".jpg") and re.search(r'Plan', filename) is not None:
            tpr.add_picture(os.getcwd() + '/{}/{}'.format(fn, filename), height=Cm(23))
        
except BaseException  as e:
    beep(4)
    easygui.msgbox('Failed to upload image "Plan": \n' + str(e), title='Error!')
    logger.error(str(e))
    sys.exit()        


# In[ ]:


# Find the place where to add The Plan picture
indx=0
for p in c.paragraphs:
    if p.text == 'Приложение №2. План квартиры':
        tpindx = indx
    indx+=1


# In[ ]:


# Move The Wall picture table to it's place
move_table_after(tpt, c.paragraphs[tpindx])


# In[ ]:





# ## Vedomost

# In[ ]:


print('Filling Defect statement...')


# In[ ]:


# Convert DataFrame to Table and add it to the Conclusion
# add a table to the end and create a reference variable
v = c.add_table(dv.shape[0]+1, dv.shape[1])
# add the header rows.
for j in range(dv.shape[-1]):
    v.cell(0,j).text = dv.columns[j]
# add the rest of the data frame and center it
for i in range(dv.shape[0]):
    for j in range(dv.shape[-1]):
        v.cell(i+1,j).text = str(dv.values[i,j])
        v.cell(i+1,j).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
# Change table style
v.style = 'Table Grid'
# Align the table
v.alignment = WD_TABLE_ALIGNMENT.CENTER        


# In[ ]:


# Avoid autofitting if the table too long
if dv.shape[1] > 10:
    v.autofit = False 
    v.allow_autofit = False


# In[ ]:


# Set table columns width
widths = [Cm(0.9), Cm(4.0), Cm(1.1), Cm(1.4)]

# Create columns widths table
widthstab = pd.DataFrame({'Room':['Кухня','Кухня-гост','Гост','Спальня', 'Комната', 'Гард', 'Кор', 'Балкон', 'Ванна', 'Сан'], 'Width':[1.3, 1.8, 1.8, 1.7, 1.7, 1.5, 1.2, 1.6, 1.5, 1.2]})


# In[ ]:


# Apply these widths to the table, vertically center it's cells and change font to Pt10 (TableStyle)
# rown = 0
for row in v.rows:
    # Apply widths for permanent columns
    for idx in range(dv.shape[1]):
        if idx < 3:
            row.cells[idx].width = widths[idx]
            row.cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            row.cells[idx].paragraphs[0].style = c.styles['TableStyle']
        elif idx == dv.shape[1]-1:
            row.cells[idx].width = widths[3]
            row.cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            row.cells[idx].paragraphs[0].style = c.styles['TableStyle']
        # Apply widths for additional columns
        else:
            for i in range(len(widthstab)):
                if dv.values[0 , idx].startswith(widthstab['Room'][i]):
                    row.cells[idx].width = Cm(widthstab['Width'][i])
                    row.cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    row.cells[idx].paragraphs[0].style = c.styles['TableStyle']


# In[ ]:


# Merge header's cells and delete strings after merging
# Vertical merge
for i in (0, 1, 2, dv.shape[1]-1):
    v.cell(0, i).merge(v.cell(1, i)).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    delete_paragraph(v.cell(0, i).paragraphs[-1])
# Horizontal merge
for j in range(3, dv.shape[1]-2):
    v.cell(0, j).merge(v.cell(0, j+1)).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    delete_paragraph(v.cell(0, j).paragraphs[-1])


# In[ ]:


# Align Names of the Jobs left
for i in range(2, dv.shape[0]+1):
    v.cell(i, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT


# In[ ]:


# Find the place where to add Vedomost table
indx=0
for p in c.paragraphs:
    if p.text == 'Дефектная ведомость':
        dvindx = indx
    indx+=1


# In[ ]:


# Move Vedomost table to it's place
move_table_after(v, c.paragraphs[dvindx])


# In[ ]:


print('  Done')


# In[ ]:





# ## Contractors

# #### Table

# In[ ]:


# Convert DataFrame to Table and add it to the Conclusion
# add a table to the end and create a reference variable
co = c.add_table(cnt.shape[0]+1, cnt.shape[1])
# add the header rows.
for j in range(cnt.shape[-1]):
    co.cell(0,j).text = cnt.columns[j]
# add the rest of the data frame and center it
for i in range(cnt.shape[0]):
    for j in range(cnt.shape[-1]):
        co.cell(i+1,j).text = str(cnt.values[i,j])
        co.cell(i+1,j).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
# Change table style
co.style = 'Table Grid'
# Align the table
co.alignment = WD_TABLE_ALIGNMENT.LEFT   


# In[ ]:


# Set table columns width
widths = [Cm(1.0), Cm(5.5)]
# Apply these widths to the table, vertically center it's cells and change font to Pt10 (TableStyle)
for row in co.rows:
    for idx, w in enumerate(widths):
        row.cells[idx].width = w
        row.cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER


# In[ ]:


# Align Names of the Contractors left
for i in range(1, cnt.shape[0]+1):
    co.cell(i, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT


# In[ ]:


# Find the place where to add Contractors table
indx=0
for p in c.paragraphs:
    if p.text.startswith('Единичные расценки были взяты из'):
        cntindx = indx
    indx+=1


# In[ ]:


# Move Contractors table to it's place
move_table_after(co, c.paragraphs[cntindx])


# In[ ]:





# #### ListLikeTable

# In[ ]:


# Convert DataFrame to ListLikeTable and add it to the Conclusion
# add a table to the end and create a reference variable
colist = c.add_table(cnl.shape[0], cnl.shape[1])
for i in range(cnl.shape[0]):
    colist.cell(i,0).text = '-'
    colist.cell(i,1).text = str(cnl.values[i,1])
    colist.cell(i,1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
# Align the table
colist.alignment = WD_TABLE_ALIGNMENT.LEFT  


# In[ ]:


# Set listliketable columns width
widths = [Cm(1.0), Cm(9.5)]
# Apply these widths to the table, vertically center it's cells and change font to Pt10 (TableStyle)
for row in colist.rows:
    for idx, w in enumerate(widths):
        row.cells[idx].width = w
        row.cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER


# In[ ]:


# Align '-' right
for i in range(cnl.shape[0]):
    colist.cell(i, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.RIGHT


# In[ ]:


# Find the place where to add Contractors listliketable
indx=0
for p in c.paragraphs:
    if p.text.startswith('Расценки взяты из следующих веб-сайтов'):
        cnllistindx = indx
    indx+=1


# In[ ]:


# Move Contractors listliketable to it's place
move_table_after(colist, c.paragraphs[cnllistindx])


# In[ ]:





# ## Prices

# In[ ]:


print('Filling cost of work and materials Tables...')


# In[ ]:


# Convert DataFrame to Table and add it to the Conclusion
# add a table to the end and create a reference variable
r = c.add_table(pr.shape[0]+1, pr.shape[1])
# add the header rows.
for j in range(pr.shape[-1]):
    r.cell(0,j).text = pr.columns[j]
# add the rest of the data frame and center it
for i in range(pr.shape[0]):
    for j in range(pr.shape[-1]):
        r.cell(i+1,j).text = str(pr.values[i,j])
        r.cell(i+1,j).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
# Change table style
r.style = 'Table Grid'
# Align the table
r.alignment = WD_TABLE_ALIGNMENT.CENTER   


# In[ ]:


# Avoid autofitting
r.autofit = False 
r.allow_autofit = False


# In[ ]:


# Set table columns width
widths = [Cm(0.9), Cm(5.7), Cm(1.1), Cm(1.4), Cm(1.2), Cm(1.2), Cm(1.2), Cm(1.2), Cm(2.0), Cm(2.1)]
# Apply these widths to the table, vertically center it's cells and change font to Pt10 (TableStyle)
for row in r.rows:
    for idx, w in enumerate(widths):
        row.cells[idx].width = w
        row.cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row.cells[idx].paragraphs[0].style = c.styles['TableStyle']


# In[ ]:


# Merge header's cells and delete strings after merging
# Vertical merge
for i in (0, 1, 2, 3, 8, 9):
    r.cell(0, i).merge(r.cell(1, i)).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    delete_paragraph(r.cell(0, i).paragraphs[-1])
# Horizontal merge
for j in range(4, 7):
    r.cell(0, j).merge(r.cell(0, j+1)).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    delete_paragraph(r.cell(0, j).paragraphs[-1])


# In[ ]:


# Align Names of the Jobs and materials left
for i in range(2, pr.shape[0]+1):
    r.cell(i, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
    # change subtotal cells stylle to bold
    if r.cell(i, 1).paragraphs[0].text.startswith('Итого') or r.cell(i, 1).paragraphs[0].text.startswith('Стоимость'):
        r.cell(i, 1).paragraphs[0].style = c.styles['SubtitleStyle']
        r.cell(i, 9).paragraphs[0].style = c.styles['SubtitleStyle']


# In[ ]:


# Repeat header rows
set_repeat_table_header(r.rows[0])


# In[ ]:


# Find the place where to add Price table
indx=0
for p in c.paragraphs:
    if p.text == 'Расчёт стоимости устранения отклонений':
        prindx = indx
    indx+=1


# In[ ]:


# Move Price table to it's place
move_table_after(r, c.paragraphs[prindx])


# In[ ]:





# ## Windows

# In[ ]:


# What if there are windows job
if wn.shape[0]>1:
    
    # Convert DataFrame to Table and add it to the Conclusion
    # add a table to the end and create a reference variable
    wi = c.add_table(wn.shape[0]+1, wn.shape[1])
    # add the header rows.
    for j in range(wn.shape[-1]):
        wi.cell(0,j).text = wn.columns[j]
    # add the rest of the data frame and center it
    for i in range(wn.shape[0]):
        for j in range(wn.shape[-1]):
            wi.cell(i+1,j).text = str(wn.values[i,j])
            wi.cell(i+1,j).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Change table style
    wi.style = 'Table Grid'
    # Align the table
    wi.alignment = WD_TABLE_ALIGNMENT.CENTER   
    
    # Avoid autofitting
    wi.autofit = False 
    wi.allow_autofit = False
    
    # Set table columns width
    widths = [Cm(0.9), Cm(7.7), Cm(1.1), Cm(1.4), Cm(1.4), Cm(1.4), Cm(2.0), Cm(2.1)]
    # Apply these widths to the table, vertically center it's cells and change font to Pt10 (TableStyle)
    for row in wi.rows:
        for idx, w in enumerate(widths):
            row.cells[idx].width = w
            row.cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            row.cells[idx].paragraphs[0].style = c.styles['TableStyle']

    # Merge header's cells and delete strings after merging
    # Vertical merge
    for i in (0, 1, 2, 3, 6, 7):
        wi.cell(0, i).merge(wi.cell(1, i)).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        delete_paragraph(wi.cell(0, i).paragraphs[-1])
    # Horizontal merge
    for j in range(4, 5):
        wi.cell(0, j).merge(wi.cell(0, j+1)).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        delete_paragraph(wi.cell(0, j).paragraphs[-1]) 
    
    # Align Names of the Jobs and materials left
    for i in range(2, wn.shape[0]+1):
        wi.cell(i, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
        # change subtotal cells stylle to bold
        if wi.cell(i, 1).paragraphs[0].text.startswith('Итого окна'):
            wi.cell(i, 1).paragraphs[0].style = c.styles['SubtitleStyle']
            wi.cell(i, 7).paragraphs[0].style = c.styles['SubtitleStyle'] 

    # Find the place where to add Windows table
    indx=0
    for p in c.paragraphs:
        if p.text.startswith('Итого все работы и материалы'):
            wnindx = indx
        indx+=1 
    
    # Add empty row to prevent concatination of tables
    c.paragraphs[wnindx-1]._p.addnext(c.add_paragraph("")._p)
    
    # Move Windows table to it's place
    move_table_after(wi, c.paragraphs[wnindx-1])


# In[ ]:





# ## Ceiling

# In[ ]:


# What if there are ceiling job
if cl.shape[0]>1:
    # Replace Contractors number according to its quantity
    cl.values[0, 4] = cn.values[cn.shape[0]-3, 0]
    cl.values[0, 5] = cn.values[cn.shape[0]-2, 0]
    cl.values[0, 6] = cn.values[cn.shape[0]-1, 0]
    
    # Convert DataFrame to Table and add it to the Conclusion
    # add a table to the end and create a reference variable
    ce = c.add_table(cl.shape[0]+1, cl.shape[1])
    # add the header rows.
    for j in range(cl.shape[-1]):
        ce.cell(0,j).text = cl.columns[j]
    # add the rest of the data frame and center it
    for i in range(cl.shape[0]):
        for j in range(cl.shape[-1]):
            ce.cell(i+1,j).text = str(cl.values[i,j])
            ce.cell(i+1,j).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Change table style
    ce.style = 'Table Grid'
    # Align the table
    ce.alignment = WD_TABLE_ALIGNMENT.CENTER   
    
    # Avoid autofitting
    ce.autofit = False 
    ce.allow_autofit = False
    
    # Set table columns width
    widths = [Cm(0.9), Cm(6.6), Cm(1.1), Cm(1.4), Cm(1.3), Cm(1.3), Cm(1.3), Cm(2.0), Cm(2.1)]
    # Apply these widths to the table, vertically center it's cells and change font to Pt10 (TableStyle)
    for row in ce.rows:
        for idx, w in enumerate(widths):
            row.cells[idx].width = w
            row.cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            row.cells[idx].paragraphs[0].style = c.styles['TableStyle']
    
    # Merge header's cells and delete strings after merging
    # Vertical merge
    for i in (0, 1, 2, 3, 7, 8):
        ce.cell(0, i).merge(ce.cell(1, i)).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        delete_paragraph(ce.cell(0, i).paragraphs[-1])
    # Horizontal merge
    for j in range(4, 6):
        ce.cell(0, j).merge(ce.cell(0, j+1)).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        delete_paragraph(ce.cell(0, j).paragraphs[-1])
    
    # Align Names of the Jobs left
    for i in range(2, cl.shape[0]+1):
        ce.cell(i, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
        # change subtotal cells stylle to bold
        if ce.cell(i, 1).paragraphs[0].text.startswith('Итого потолок'):
            ce.cell(i, 1).paragraphs[0].style = c.styles['SubtitleStyle']
            ce.cell(i, 8).paragraphs[0].style = c.styles['SubtitleStyle']
    
    # Find the place where to add Ceiling table
    indx=0
    for p in c.paragraphs:
        if p.text.startswith('Итого все работы и материалы'):
            clindx = indx
        indx+=1
    
    # Add empty row to prevent concatination of tables
    c.paragraphs[clindx-1]._p.addnext(c.add_paragraph("")._p)
    
    # Move Ceiling table to it's place
    move_table_after(ce, c.paragraphs[clindx-1])    


# In[ ]:


print('  Done')


# In[ ]:





# ## Photos

# In[ ]:


print('Uploading Photos...')


# In[ ]:


try:
    # Read PDF file
    for pdf_file in os.listdir(os.getcwd() + '/{}/'.format(fn)):
#         if pdf_file.endswith(".pdf") and not pdf_file.startswith('Conclusion'):
        if pdf_file.endswith(".pdf") and pdf_file.startswith('report'):    
            # Convert it to JPG
            pages = convert_from_path(os.getcwd() + '/{}/{}'.format(fn, pdf_file), 300, poppler_path = os.getcwd() + '/SYSTEM/Poppler/poppler-21.08.0/Library/bin')
            # Create table where to put Photos
            tft = c.add_table(len(pages), 1)
#             # Cancel autofitting
#             tft.autofit = False 
#             tft.allow_autofit = False
            # Add each page to the table
            for page in pages:         
                # Save the page
                page.save(os.getcwd() + '/{}/page.jpg'.format(fn))            
                # Add paragraph to the cell and run it
                tfp = tft.rows[pages.index(page)].cells[0].add_paragraph()
                tfr = tfp.add_run()
                # Add picture to this paragraph
                 # Define picture height
                if pages.index(page) in [0, len(pages)-1]:
                    h = 25.5
                else:
                    h = 27
                tfr.add_picture(os.getcwd() + '/{}/page.jpg'.format(fn), height=Cm(h))
                 # Delete empty string before picture
                delete_paragraph(tft.rows[pages.index(page)].cells[0].paragraphs[-2])
    # Delete the page file
    os.remove(os.getcwd() + '/{}/page.jpg'.format(fn))

except BaseException  as e:
    beep(4)
    easygui.msgbox('Failed to upload inspection photos: \n' + str(e), title='Error!')
    logger.error(str(e))
    sys.exit()


# In[ ]:


# Find the place where to add Photos
indx=0
for p in c.paragraphs:
    if p.text.endswith('Фотографии, выполненные экспертом во время осмотра'):
        tfindx = indx
    indx+=1


# In[ ]:


# Move Phtos table to it's place
move_table_after(tft, c.paragraphs[tfindx])


# In[ ]:


print('  Done')


# In[ ]:





# ## Certificates

# In[ ]:


print('Uploading Annexes...')


# In[ ]:


# Convert DataFrame to ListLikeTable and add it to the Conclusion
# add a table to the end and create a reference variable
crlist = c.add_table(ct.shape[0], 1)
for i in range(ct.shape[0]):
    crlist.cell(i,0).text = str(ct.values[i])
    crlist.cell(i,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
# Align the table
crlist.alignment = WD_TABLE_ALIGNMENT.LEFT  


# In[ ]:


# Find the place where to add Certificates table
indx=0
for p in c.paragraphs:
    if p.text.startswith('Оборудование, использованное при производстве'):
        ctindx = indx
    indx+=1


# In[ ]:


# Move Certificates table to it's place
move_table_after(crlist, c.paragraphs[ctindx])


# In[ ]:





# In[ ]:


# Remove NaNs at ctd
ctd = ctd.dropna()
ctd = ctd.reset_index(drop=True)


# In[ ]:


# Delete '/' at documents names
ctd = ctd.str.replace('/','')


# In[ ]:


# Find all the necessary certificates:
certs = np.array([])
for i in range(len(ctd)):
    for filename in os.listdir(os.getcwd() + '/Source/Verifications'):
        if filename.startswith(ctd.values[i]):
            #Add them to the list
            certs = np.append(certs, ctd.values[i])


# In[ ]:


# Make elements quantity even:
if len(certs) % 2 != 0:
    certs = np.append(certs, '')


# In[ ]:


# Reshape array to two columns array:
certs = np.reshape(certs, (-1, 2))


# In[ ]:


# Create the table where to put The Walls picture in
tct = c.add_table(certs.shape[0], 2)


# In[ ]:


try:
    # Insert Certificates to the table
    for i in range(certs.shape[0]):
        for j in range(2):
            # Add paragraph to the cell and run it
            tcp = tct.rows[i].cells[j].add_paragraph()
            tcr = tcp.add_run()
            # Add picture to this paragraph
            if certs[i,j] != '':
                certname = f'{certs[i,j]}'
                tcr.add_picture(os.getcwd() + '/Source/Verifications/{}.jpg'.format(certname), height=Cm(11))

except BaseException  as e:
    beep(4)
    easygui.msgbox('Failed to upload certificates: \n' + str(e), title='Error!')
    logger.error(str(e))
    sys.exit()


# In[ ]:


# Find the place where to add Certificates
indx=0
for p in c.paragraphs:
    if p.text.startswith('Приложение №4. Поверки '):
        certindx = indx
    indx+=1


# In[ ]:


# Move Certificates to it's place
move_table_after(tct, c.paragraphs[certindx])


# In[ ]:





# ## Invoices

# In[ ]:


# Find Invoices start paragraphs
indx=0
for p in c.paragraphs:
    if p.text.startswith('Приложение №5. Счета'):
        invindx = indx
    indx+=1


# In[ ]:


try:
    # What if there are no invoices (no windows or no standard windows)
    if standwind == 0 or wn.shape[0]<=1:
        # Delete the paragraph and empty string
        delete_paragraph(c.paragraphs[invindx-1])
        delete_paragraph(c.paragraphs[invindx-1])
    # What if there are invoices (windows)
    else:
        # Create table to put Invoices in
        tit = c.add_table(2, 1)
        # Insert Invoices to the table
        # Iterate trough the Invoices
        i = 0
        for inv_file in os.listdir(os.getcwd() + '/Source/Invoices'):
            i = i+1
    #         invname = f'{inv_file}'
            # Add paragraph to the cell and run it
            tip = tit.rows[i-1].cells[0].add_paragraph()
            tir = tip.add_run()
            # Add picture to this paragraph
            tir.add_picture(os.getcwd() + '/Source/Invoices/{}'.format(inv_file), height=Cm(24))
            
except BaseException  as e:
    beep(4)
    easygui.msgbox('Failed to upload invoices: \n' + str(e), title='Error!')
    logger.error(str(e))
    sys.exit()    


# In[ ]:


print('  Done')


# In[ ]:





# ## Final edit and save

# In[ ]:


print('Saving the Сonclusion...')


# In[ ]:


# Put the data into Conclusion template
context = {'contract' : ag['Значение'][0],
           'date' : ag['Значение'][1],
           'start_date' : ag['Значение'][2],
           'end_date' : ag['Значение'][3],
           'nstart_time' : ag['Значение'][10],
           'nend_time' : ag['Значение'][11],
           'start_time' : resstart,
           'end_time' : resend,
           'sum' : ag['Значение'][17],
           'sum_words' : num2text(float(ag['Значение'][17])),
           'address' : ag['Значение'][4],
           'client' : ag['Значение'][7],    
           'interesant' : interesant,
           'temperature_inside' : ag['Значение'][12],
           'humidity' : ag['Значение'][13],
           'hygrometer' : ag['Значение'][14],
           'temperature_outside' : ostemp}
c.render(context)


# In[ ]:


# Update table of content
namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
# add child to doc.settings element
element_updatefields = lxml.etree.SubElement(c.settings.element, f"{namespace}updateFields")
element_updatefields.set(f"{namespace}val", "true")


# In[ ]:





# In[ ]:


# Save Conclusion
c.save(os.getcwd() + '/{}/{}.docx'.format(fn, 'Conclusion '+ agnr))


# In[ ]:


# Convert conclusion to pdf
wdFormatPDF = 17
inputFile = os.path.abspath(os.getcwd() + '/{}/{}.docx'.format(fn, 'Conclusion '+ agnr))
outputFile = os.path.abspath(os.getcwd() + '/{}/{}.pdf'.format(fn, 'Conclusion '+ agnr))
word = win32com.client.Dispatch('Word.Application')
doc = word.Documents.Open(inputFile)
doc.SaveAs(outputFile, FileFormat=wdFormatPDF)
doc.Close(-1)


# In[ ]:





# In[ ]:


# Rename initial Smeta file if it is Creation
if scrp == 'creation':
    os.rename(os.getcwd() + '/{}/Estimate.xlsx'.format(fn), os.getcwd() + '/{}/{}.xlsx'.format(fn, 'Estimate '+ apnr))
# Delete initial Smeta file if it is Correction
if scrp == 'correction':
    try:
        os.remove(os.getcwd() + '/{}/Estimate.xlsx'.format(fn))
    except:
        print('no such file')


# In[ ]:





# In[ ]:


# Copy all the files to a new folder
 # Remove the same folder if it exists
if os.path.isdir(os.getcwd() + '/Сonclusions/{}'.format(name + ' ' + agnr)):
    beep(4)
    if easygui.buttonbox('A folder with the same name already exists, do you want to replace it?', 'Attention!!!', ('Yes', 'No, leave all files in the "Work" folder')) == 'Yes':
        shutil.rmtree(os.getcwd() + '/Сonclusions/{}'.format(name + ' ' + agnr))
        
        # Copy files to a new folder
        shutil.copytree(os.getcwd() + '/{}'.format(fn), os.getcwd() + '/Сonclusions/{}'.format(name + ' ' + agnr))
        
        # Delete all the files from the main working folder
        while len(os.listdir(os.getcwd() + '/{}'.format(fn))) > 0:
            try:
                for files in os.listdir(os.getcwd() + '/{}'.format(fn)):
                    os.remove(os.getcwd() + '/{}/{}'.format(fn, files))
            except BaseException  as e:
                beep(4)
                easygui.msgbox('Failed to clean folder "Work". Close open files.: \n' + str(e), title='Error!')
else:
    shutil.copytree(os.getcwd() + '/{}'.format(fn), os.getcwd() + '/Сonclusions/{}'.format(name + ' ' + agnr))

    # Delete all the files from the main working folder
    while len(os.listdir(os.getcwd() + '/{}'.format(fn))) > 0:
        try:
            for files in os.listdir(os.getcwd() + '/{}'.format(fn)):
                os.remove(os.getcwd() + '/{}/{}'.format(fn, files))
        except BaseException  as e:
            beep(4)
            easygui.msgbox('Failed to clean folder "Work". Close open files.: \n' + str(e), title='Error!')


# In[ ]:


# Copy Smeta again to the main folder
try:
    # Copy initial file back to working folder
    shutil.copy(os.getcwd() + '/Source/Estimate.xlsx', os.getcwd() + '/{}'.format(fn))

except BaseException  as e:
    beep(4)
    easygui.msgbox('Failed to upload template "Estimate": \n' + str(e), title='Error!')
    logger.error(str(e))
    sys.exit()


# In[ ]:


# Add sound
if ag['Unnamed: 4'][4] > 0:
    beep(ag['Unnamed: 4'][4])


# In[ ]:


print('  Conclusion for ' + name + ' has been drawn up!')


# In[ ]:


time.sleep(3)

